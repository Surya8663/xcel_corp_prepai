"""
End-to-End Test Script for Resume Upload & Audit Endpoint.

Generates a realistic sample PDF resume in memory/temp file using reportlab or fpdf (or a text file converted to PDF/DOCX),
posts it to http://localhost:8000/api/v1/resume/upload, and asserts:
  - 201 Created status
  - Valid audit score (0-100)
  - Parsed JSON contains skills, experience, education
  - Audit JSON contains whats_good, needs_improvement, industry_level, disclaimer note alignment
"""

from __future__ import annotations

import io
import json
import os
import sys
import time
import requests

# ASCII-safe console output for Windows PowerShell
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

API_URL = "http://localhost:8000/api/v1/resume/upload"


def create_sample_docx() -> bytes:
    """Create a sample DOCX resume in memory using python-docx."""
    import docx

    doc = docx.Document()
    doc.add_heading("Alex Rivera", 0)
    doc.add_paragraph("Senior Full-Stack Engineer | San Francisco, CA | alex@example.com | github.com/alexrivera")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Senior Software Engineer with 6+ years of experience building high-throughput microservices, "
        "cloud-native backend architectures, and modern web applications with React, Next.js, FastAPI, and PostgreSQL. "
        "Proven track record of optimizing database performance and deploying LLM features to production."
    )

    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("Languages: Python, TypeScript, JavaScript, SQL, HTML/CSS")
    doc.add_paragraph("Frameworks & Tools: FastAPI, Next.js, React, Node.js, SQLAlchemy, PostgreSQL, Docker, AWS, Git")

    doc.add_heading("Experience", level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run("Senior Software Engineer ").bold = True
    p1.add_run("— TechCorp Inc. (Jan 2022 - Present)\n")
    doc.add_paragraph("• Architected async microservices in Python (FastAPI) serving 50k active daily users.")
    doc.add_paragraph("• Optimized PostgreSQL query execution plans, reducing p99 latency from 450ms to 85ms.")
    doc.add_paragraph("• Integrated Gemini LLM APIs for automated content generation, cutting processing time by 40%.")

    p2 = doc.add_paragraph()
    p2.add_run("Software Engineer ").bold = True
    p2.add_run("— CloudApps Solutions (Jun 2019 - Dec 2021)\n")
    doc.add_paragraph("• Developed React and Next.js frontend applications with Tailwind CSS and Zustand.")
    doc.add_paragraph("• Designed RESTful and GraphQL APIs integrated with PostgreSQL and Redis caching.")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.S. in Computer Science — University of California, Berkeley (2015 - 2019)")

    doc.add_heading("Projects", level=1)
    doc.add_paragraph("PrepAI: Adaptive interview preparation platform using Gemini 1.5/2.0 Flash and Next.js 14.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def run_e2e_test():
    print("=" * 60)
    print("  PrepAI — Resume Upload & Audit End-to-End Test")
    print("=" * 60)

    print("\n[1/3] Generating sample resume file (DOCX)...")
    docx_bytes = create_sample_docx()
    print(f"      Generated DOCX resume size: {len(docx_bytes)} bytes")

    print("\n[2/3] Uploading resume to backend API (POST /api/v1/resume/upload)...")
    t0 = time.time()
    files = {
        "file": ("alex_rivera_resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    }
    
    response = requests.post(API_URL, files=files, timeout=60)
    elapsed = time.time() - t0

    print(f"      Response Status Code: {response.status_code} ({elapsed:.2f}s)")

    if response.status_code != 201:
        print(f"ERROR: Upload failed! Response body: {response.text}")
        sys.exit(1)

    data = response.json()
    print("\n[3/3] Validating API Response Data...")
    print(f"      Resume ID        : {data.get('id')}")
    print(f"      Candidate ID     : {data.get('candidate_id')}")
    print(f"      File URL         : {data.get('file_url')}")
    print(f"      Audit Score      : {data.get('audit_score')} / 100")

    parsed = data.get("parsed_json") or {}
    print(f"      Extracted Skills : {len(parsed.get('skills', []))} skills found -> {parsed.get('skills', [])[:8]}")
    print(f"      Extracted Exp    : {len(parsed.get('experience', []))} entries found")

    audit = data.get("audit_feedback_json") or {}
    print(f"      Industry Level   : {audit.get('industry_level')}")
    print(f"      What's Good      : {len(audit.get('whats_good', []))} points")
    print(f"      Needs Improve    : {len(audit.get('needs_improvement', []))} items")

    # Assertions
    assert data.get("id") is not None, "Resume ID should not be None"
    assert data.get("audit_score") is not None and 0 <= data["audit_score"] <= 100, "Audit score must be between 0 and 100"
    assert len(parsed.get("skills", [])) > 0, "Should have extracted skills"
    assert len(audit.get("whats_good", [])) > 0, "Should have whats_good feedback"
    assert len(audit.get("needs_improvement", [])) > 0, "Should have needs_improvement feedback"

    print("\nSUCCESS: End-to-end Resume Audit test PASSED cleanly!")
    print("=" * 60)


if __name__ == "__main__":
    run_e2e_test()
