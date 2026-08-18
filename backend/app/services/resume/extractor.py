"""
PrepAI — Resume text extractor.

Extracts raw text from uploaded PDF or DOCX files using real parsing libraries.
No mocked or hardcoded text — extraction is always from the actual file bytes.
"""

from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB


class ExtractionError(Exception):
    """Raised when text extraction from a file fails."""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract plain text from a PDF using pdfplumber.
    Preserves page-level structure with line breaks between pages.
    """
    try:
        import pdfplumber
    except ImportError as exc:
        raise ExtractionError("pdfplumber is not installed. Run: pip install pdfplumber") from exc

    try:
        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    pages.append(text)
                else:
                    logger.warning("PDF page %d returned no text (may be image-based)", i + 1)

        extracted = "\n\n".join(pages).strip()
        if not extracted:
            raise ExtractionError(
                "PDF appears to be image-based or has no selectable text. "
                "Please upload a text-based PDF."
            )
        logger.info("Extracted %d characters from PDF (%d pages)", len(extracted), len(pages))
        return extracted

    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"Failed to parse PDF: {exc}") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract plain text from a DOCX file using python-docx.
    Reads both paragraph text and table cell text.
    """
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError("python-docx is not installed. Run: pip install python-docx") from exc

    try:
        doc = docx.Document(io.BytesIO(file_bytes))

        parts: list[str] = []

        # Main body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Table cells (resumes often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in parts:
                        parts.append(text)

        extracted = "\n".join(parts).strip()
        if not extracted:
            raise ExtractionError(
                "DOCX file appears to be empty or contains no readable text paragraphs."
            )
        logger.info("Extracted %d characters from DOCX", len(extracted))
        return extracted

    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"Failed to parse DOCX: {exc}") from exc


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Dispatch to the correct extractor based on file extension.

    Args:
        filename: Original filename including extension.
        file_bytes: Raw file content as bytes.

    Returns:
        Extracted plain text string.

    Raises:
        ExtractionError: If the file type is unsupported or parsing fails.
    """
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ExtractionError(
            f"File is too large ({len(file_bytes) // 1024 // 1024} MB). "
            f"Maximum allowed size is {MAX_FILE_SIZE_BYTES // 1024 // 1024} MB."
        )

    ext = "." + filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ExtractionError(
            f"Unsupported file type: '{ext or filename}'. "
            f"Please upload a PDF or DOCX file."
        )
